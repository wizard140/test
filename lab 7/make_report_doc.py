from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def set_margins(section):
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)


def set_font(run, name="Aptos", size=11, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    set_font(run, size=16 - 2 * (level - 1), bold=True)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    run = p.add_run(text)
    set_font(run)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_font(run)
        p.paragraph_format.space_after = Pt(2)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


doc = Document()
for section in doc.sections:
    set_margins(section)

style = doc.styles["Normal"]
style.font.name = "Aptos"
style._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
style.font.size = Pt(11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("CUDA Convolution Lab Report Template")
set_font(r, size=18, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Prepared from the provided assignment prompt and the attached CUDA solution package")
set_font(r, size=11)

add_body(doc, "This report template is written to match the code package in this folder. It already covers the design, hypotheses, filters, testing method, and analysis workflow. After you run the program on your CUDA machine, replace the bracketed placeholders with your measured values, figures, and short discussion paragraphs.")

add_heading(doc, "1. Project Overview", 1)
add_body(doc, "This project studies how different CUDA optimization strategies affect 2D grayscale image convolution. The program applies the same filters to the same synthetic test images using five implementations: a single-threaded CPU baseline, a tuned but otherwise plain CUDA kernel, a tiled shared-memory kernel, a constant-memory kernel, and a tiled kernel that also uses constant memory for the filter.")
add_body(doc, "Input: a grayscale image represented as a 2D float array and a 2D convolution filter. Output: a same-size grayscale output image produced using zero-padding at the borders. Constraints: the image and filter must fit in available memory, and the constant-memory versions are limited by the constant-filter array in the code, which is currently sized for filters up to 31 x 31.")
add_body(doc, "Compared with a minimal baseline implementation, this version uses pinned host memory to reduce page-fault overhead during transfers, tuned launch dimensions for the plain GPU kernel, shared-memory tiling for data reuse, and constant memory for cached filter reads.")

add_heading(doc, "2. Problem Formulation and Requirements", 1)
add_body(doc, "For each output pixel (x, y), convolution computes the weighted sum of neighboring input pixels under a selected filter mask. The CPU and GPU versions both use the same zero-padded convolution rule so that their outputs can be compared directly.")
add_bullets(doc, [
    "The CPU version must remain single-threaded and serves as the baseline.",
    "The GPU timing must include host-to-device and device-to-host memory transfers.",
    "Verification time must not be included in the timed measurements.",
    "The same images and filters must be used across every implementation.",
    "The two provided 3 x 3 filters and two user-chosen 5 x 5 filters must all be tested.",
])

add_heading(doc, "3. Implementations", 1)
add_body(doc, "The code package implements the following five versions.")

table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
headers = ["Version", "What it does", "Why it may be faster"]
for i, text in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = text
    shade_cell(cell, "D9EAF7")

rows = [
    ("cpu_single", "Single-threaded nested-loop convolution on the CPU.", "No GPU overhead, but no parallelism."),
    ("gpu_global", "One thread per output element with tuned grid/block sizes, reading image and filter from global memory.", "Massive parallelism even without data reuse optimizations."),
    ("gpu_tiled", "Loads an input tile plus halo into shared memory before computing outputs.", "Reduces repeated global-memory reads for neighboring pixels."),
    ("gpu_constant", "Reads the filter from constant memory while still reading the image from global memory.", "All threads reuse the same filter coefficients, so cached constant-memory reads can help."),
    ("gpu_tiled_constant", "Combines shared-memory tiling for the image with constant-memory filter reads.", "Gets both image data reuse and cheap repeated filter reads."),
]
for row in rows:
    cells = table.add_row().cells
    for i, text in enumerate(row):
        cells[i].text = text

add_heading(doc, "4. Hypotheses and Predictions", 1)
add_body(doc, "Before running the experiments, the expected ordering from slowest to fastest on medium and large images is cpu_single, gpu_global, gpu_constant, gpu_tiled, and gpu_tiled_constant. The CPU version should only be competitive on very small images because it avoids transfer overhead. Once the image becomes large enough, the GPU should pull ahead because thousands of pixels can be processed in parallel.")
add_bullets(doc, [
    "The plain global-memory kernel should beat the CPU once transfer overhead is amortized.",
    "The tiled kernel should beat the plain global-memory kernel because neighboring pixels are reused inside each block.",
    "The constant-memory kernel should improve over the plain global-memory kernel because every thread reads the same filter values.",
    "The tiled + constant-memory kernel is expected to be the fastest overall because it reduces both image-memory traffic and filter-memory traffic.",
    "The 5 x 5 filters should take measurably longer than the 3 x 3 filters because each output requires more multiply-add work and a larger halo in the tiled kernels.",
])

add_heading(doc, "5. Filters Used", 1)
add_body(doc, "The program tests the following four filters.")
filter_table = doc.add_table(rows=1, cols=3)
filter_table.style = "Table Grid"
for i, text in enumerate(["Filter", "Size", "Effect"]):
    cell = filter_table.rows[0].cells[i]
    cell.text = text
    shade_cell(cell, "D9EAF7")

filter_rows = [
    ("emboss_3x3", "3 x 3", "Highlights diagonal intensity changes and gives an embossed / directional-edge look."),
    ("sharpen_3x3", "3 x 3", "Emphasizes local contrast by boosting the center pixel relative to its immediate neighbors."),
    ("gaussian_5x5", "5 x 5", "Applies a weighted blur that smooths noise while preserving structure better than a uniform blur."),
    ("edge_5x5", "5 x 5", "Acts like a stronger high-pass / edge-emphasis filter over a larger neighborhood."),
]
for row in filter_rows:
    cells = filter_table.add_row().cells
    for i, text in enumerate(row):
        cells[i].text = text

add_heading(doc, "6. Test Images", 1)
add_body(doc, "The package uses synthetic grayscale images generated directly in memory so that file I/O is not part of the timed section. This satisfies the assignment note that synthetic images may be used for testing.")
add_bullets(doc, [
    "gradient_512: a smooth gradient used to show broad intensity changes.",
    "checker_1024: a sharp-edged checkerboard pattern that makes blur and edge effects easy to see.",
    "rings_2048: a circular wave pattern that shows how filters respond to repeating edges and curves.",
    "noise_3072: pseudo-random noise that is useful for testing smoothing and edge emphasis on a busy input.",
])

add_heading(doc, "7. Experimental Method", 1)
add_body(doc, "Each image/filter pair is timed for every implementation across multiple repeats. The timed GPU region starts immediately before copying the input image to the device and ends after copying the finished output image back to host memory. File I/O and output verification are performed outside the timed section. After the timed runs finish, each GPU result is compared with the CPU result to confirm consistency.")
add_bullets(doc, [
    "Build: `make`",
    "Generate timing data: `make test`",
    "Generate plots and summary files: `python3 analyze_results.py lab_output/results.csv lab_output`",
    "Optional one-step workflow: `make analyze`",
])

add_heading(doc, "8. Results", 1)
add_body(doc, "Insert the generated plots and fill the short discussion paragraphs below after running the program.")
add_body(doc, "Observed speedups over the CPU baseline: [Insert 2-4 sentences summarizing which GPU version was fastest and how speedup changed as image size increased.]")
add_body(doc, "Difference between versions: [Use the means, standard deviations, and stat_tests.txt output to explain whether the implementations differ in a statistically meaningful way.]")
add_body(doc, "Break-even point: [Use break_even_summary.csv and the speedup plot to describe the image size where GPU execution begins to outperform the CPU baseline.]")
add_body(doc, "Hypothesis check: [State whether the results supported or refuted the original predictions.]")
add_body(doc, "Filter-size effect: [State whether 5 x 5 filters were measurably slower than 3 x 3 filters and quantify the change.]")

add_heading(doc, "9. Required Figures and Tables", 1)
add_bullets(doc, [
    "Figure 1. speedup_vs_pixels.png",
    "Figure 2. runtime_vs_pixels.png",
    "Figure 3. filter_size_comparison.png",
    "Table 1. timing_summary.csv (mean and standard deviation)",
    "Table 2. speedup_summary.csv",
    "Table 3. break_even_summary.csv",
    "Appendix image examples: sample_*.pgm converted to viewable images if needed",
])

add_heading(doc, "10. Direct Answers to the Prompt Questions", 1)
answers = [
    "What speedups were observed for the multi-threaded versions over the single-threaded version? [Answer after running.]",
    "Is there a difference between the versions? [Use timing_summary.csv and stat_tests.txt.]",
    "Based on your models what size is the break-even point for the GPU versions over the CPU version? [Use break_even_summary.csv.]",
    "Does the data support or refute your hypotheses? [Answer after running.]",
    "What trends did you find? Be quantitative. [Answer after running.]",
    "Show examples of the transformed images for each filter. [Insert sample images.]",
    "Where did your test images come from? The images were generated synthetically in memory by the program, which is allowed by the prompt.",
    "What do the two 5 x 5 filters do? The Gaussian filter performs weighted smoothing, while the 5 x 5 edge filter emphasizes larger-scale intensity changes and edges.",
    "What do the two 3 x 3 filters do? The emboss filter emphasizes directional change and gives a raised-texture look, while the sharpen filter increases local contrast and crispness.",
    "Does filter size affect runtime in a measurable way? [Answer after running, using the filter size comparison plot and summary table.]",
]
add_bullets(doc, answers)

add_heading(doc, "11. Notes on Style and Code Structure", 1)
add_body(doc, "The CUDA code was written in a plain function-based style to stay close to the structure used in your earlier code: helper functions are defined above main, naming stays direct, and the timing wrappers are separated so the report can clearly explain what each measured region includes.")

path = "/mnt/data/convolution_lab_package/Convolution_Lab_Report_Template.docx"
doc.save(path)
print(path)
